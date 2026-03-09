"""Markdown-to-Word converter with professional formatting.

Converts the analysis markdown report into a styled .docx document
with proper headings, tables, bullet lists, and brand colours.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# -- Brand defaults (mirror pptx_export.py) -----------------------------------
BRAND_TITLE_COLOR = RGBColor(0x00, 0x3B, 0x70)    # Dark blue
BRAND_ACCENT_COLOR = RGBColor(0x00, 0x6B, 0xA6)   # Citi blue
BRAND_BODY_COLOR = RGBColor(0x33, 0x33, 0x33)      # Dark grey
BRAND_FONT = "Calibri"

# Table cell shading for header row
_TABLE_HEADER_FILL = "003B70"
_TABLE_HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
_TABLE_ALT_ROW_FILL = "F2F6FA"


# -----------------------------------------------------------------------------
# Style helpers
# -----------------------------------------------------------------------------

def _apply_brand_styles(doc: Document) -> None:
    """Configure document styles to match brand defaults."""
    style = doc.styles["Normal"]
    style.font.name = BRAND_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_BODY_COLOR
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        name = f"Heading {level}"
        if name in doc.styles:
            hs = doc.styles[name]
            hs.font.name = BRAND_FONT
            hs.font.color.rgb = BRAND_TITLE_COLOR if level <= 2 else BRAND_ACCENT_COLOR
            hs.font.bold = True
            hs.font.size = Pt({1: 22, 2: 16, 3: 13, 4: 12}[level])
            hs.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
            hs.paragraph_format.space_after = Pt(6)


def _shade_cell(cell: Any, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): hex_color, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_elem)


def _style_table(table: Any) -> None:
    """Apply professional styling to a table: header row shading, alternating rows, borders."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    if table.rows:
        for cell in table.rows[0].cells:
            _shade_cell(cell, _TABLE_HEADER_FILL)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = _TABLE_HEADER_FONT_COLOR
                    run.font.size = Pt(10)

    # Data rows: alternating fill, consistent font
    for idx, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            if idx % 2 == 0:
                _shade_cell(cell, _TABLE_ALT_ROW_FILL)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = BRAND_FONT


# -----------------------------------------------------------------------------
# Markdown parsing helpers
# -----------------------------------------------------------------------------

_RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_RE_BULLET = re.compile(r"^(\s*)[-*]\s+(.+)$")
_RE_NUMBERED = re.compile(r"^(\s*)\d+\.\s+(.+)$")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")
_RE_HR = re.compile(r"^-{3,}$|^\*{3,}$|^_{3,}$")


def _add_rich_text(paragraph: Any, text: str) -> None:
    """Add text with inline bold/italic formatting to a paragraph."""
    # Split by bold markers first, then italic within each segment
    parts = _RE_BOLD.split(text)
    is_bold = False
    for part in parts:
        if not part:
            is_bold = not is_bold
            continue
        if is_bold:
            run = paragraph.add_run(part)
            run.bold = True
            is_bold = False
        else:
            # Check for italic within non-bold text
            italic_parts = _RE_ITALIC.split(part)
            is_italic = False
            for ip in italic_parts:
                if not ip:
                    is_italic = not is_italic
                    continue
                if is_italic:
                    run = paragraph.add_run(ip)
                    run.italic = True
                    is_italic = False
                else:
                    paragraph.add_run(ip)
            is_bold = True  # next segment is bold


def _parse_table_block(lines: list[str]) -> list[list[str]]:
    """Parse consecutive markdown table lines into a list of rows (list of cell strings)."""
    rows: list[list[str]] = []
    for line in lines:
        if _RE_TABLE_SEP.match(line.strip()):
            continue
        match = _RE_TABLE_ROW.match(line.strip())
        if match:
            cells = [c.strip() for c in match.group(1).split("|")]
            rows.append(cells)
    return rows


# -----------------------------------------------------------------------------
# Main converter
# -----------------------------------------------------------------------------

def markdown_to_docx(markdown: str, output_path: str) -> str:
    """Convert markdown text to a professionally formatted Word document.

    Supports: headings (H1-H4), bullet lists, numbered lists, bold/italic,
    markdown tables, and horizontal rules (as section breaks).

    Args:
        markdown: The markdown content to convert.
        output_path: Absolute path for the output .docx file.

    Returns:
        Absolute path to the generated .docx file.
    """
    # Strip PPTX slide hints (<!-- SLIDE: ... -->) and other HTML comments
    markdown = re.sub(r"<!--.*?-->", "", markdown, flags=re.DOTALL)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    _apply_brand_styles(doc)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule → visual section break
        if _RE_HR.match(stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            # Add a thin horizontal line via bottom border
            pPr = para._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "1",
                    qn("w:color"): "CCCCCC",
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = _RE_HEADING.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table block (starts with |)
        if _RE_TABLE_ROW.match(stripped):
            table_lines: list[str] = []
            while i < len(lines) and _RE_TABLE_ROW.match(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table_block(table_lines)
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < n_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ""
                            para = cell.paragraphs[0]
                            _add_rich_text(para, cell_text)
                _style_table(table)
                doc.add_paragraph()  # spacing after table
            continue

        # Bullet lists
        bullet_match = _RE_BULLET.match(line)
        if bullet_match:
            text = bullet_match.group(2)
            indent = len(bullet_match.group(1))
            style = "List Bullet 2" if indent >= 2 and "List Bullet 2" in doc.styles else "List Bullet"
            para = doc.add_paragraph(style=style)
            _add_rich_text(para, text)
            i += 1
            continue

        # Numbered lists
        num_match = _RE_NUMBERED.match(line)
        if num_match:
            text = num_match.group(2)
            para = doc.add_paragraph(style="List Number")
            _add_rich_text(para, text)
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _add_rich_text(para, stripped)
        i += 1

    # Footer
    footer_section = doc.sections[-1]
    footer = footer_section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Generated by AgenticAnalytics")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = BRAND_FONT

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return str(Path(output_path).resolve())
